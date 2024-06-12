import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import tempfile

# Initialize session state for user stories if it doesn't exist
if 'user_stories' not in st.session_state:
    st.session_state.user_stories = []

# Function to create a Word document for multiple user stories
def create_user_story_doc(stories):
    doc = Document()
    doc.add_heading('User Stories', 0)

    for story in stories:
        title, priority, user_story_desc, as_a, i_want_to, functionality, so_that = story

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        table.cell(0, 0).text = f"Title: {title}"
        table.cell(0, 1).text = f"Priority: {priority}"

        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 0).text = user_story_desc

        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 0).text = f"As a {as_a}: I want to {i_want_to}"

        table.cell(3, 0).merge(table.cell(3, 1))
        functionality_cell = table.cell(3, 0)
        functionality_cell.text = "Functionality:"
        for item in functionality:
            p = functionality_cell.add_paragraph()
            p.text = item
            p.style = 'List Bullet'

        table.cell(4, 0).merge(table.cell(4, 1))
        table.cell(4, 0).text = f"So that: {so_that}"

        doc.add_paragraph()  # Add a blank line between stories

    return doc

st.title('User Story Generator')

title = st.text_input('Title', 'Design the Product Recommendation Card')
priority = st.text_input('Priority', '')
user_story_desc = st.text_input('User Story Description', 'User Story 1')
as_a = st.text_input('As a', 'Designer')
i_want_to = st.text_area('I want to', 'create a visually appealing and user-friendly product recommendation interface')
functionality = st.text_area('Functionality', 'Utilize elements from the "product information card" design, such as product image, name, and description.\nShow recommended products based on user preferences, browsing history, or other relevant data.\nMake sure there is sync in visual style and interaction pattern between the product information card and product recommendation card')
so_that = st.text_area('So that', 'all users see product suggestions without needing to refresh the page')

if st.button('Add Story'):
    functionality_list = functionality.split('\n')
    st.session_state.user_stories.append((title, priority, user_story_desc, as_a, i_want_to, functionality_list, so_that))
    st.success('User story added successfully.')

if st.button('Download All Stories'):
    if st.session_state.user_stories:
        doc = create_user_story_doc(st.session_state.user_stories)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        with open(tmp_path, 'rb') as f:
            st.download_button('Download User Stories Document', f, file_name='User_Stories.docx')
    else:
        st.warning('No user stories to download.')

if st.button('Clear All Stories'):
    st.session_state.user_stories.clear()
    st.success('All user stories cleared.')
