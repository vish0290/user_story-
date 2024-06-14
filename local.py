import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import tempfile



def doc_create(data):
    doc = Document()
    count = 1
    for stories in data:
        doc.add_heading(stories['Page_title'],0)
        for story in stories['Data']:
            table = doc.add_table(rows=5,cols=2)
            table.style = 'Table Grid'
            
            table.cell(0, 0).text = f"Title: {story['Title']}"
            table.cell(0, 1).text = f"Priority: {story['Priority']}"

            table.cell(1, 0).merge(table.cell(1, 1))
            table.cell(1, 0).text = f'User Story {count}'
            count+=1

            table.cell(2, 0).merge(table.cell(2, 1))
            table.cell(2, 0).text = story['Requirements']

            table.cell(3, 0).merge(table.cell(3, 1))
            table.cell(3, 0).text = f"Functionality: {story['Functionality']}"
            
            
            table.cell(4, 0).merge(table.cell(4, 1))
            table.cell(4, 0).text = f"So that: {story['So That']}"
            doc.add_paragraph()
            
    return doc


path = 'final.json'

with open(path,'r') as f:
    user_stories = json.load(f)

doc = doc_create(user_stories)
doc.save('stories.docx')